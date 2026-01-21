#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Convert Markdown technical document to Word format with Mermaid flowcharts as images
"""

import re
import requests
import base64
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def encode_mermaid(mermaid_code):
    """Encode Mermaid code for mermaid.ink API"""
    return base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')

def mermaid_to_image(mermaid_code, output_file):
    """Convert Mermaid code to image using mermaid.ink API"""
    try:
        # Remove leading/trailing whitespace
        mermaid_code = mermaid_code.strip()

        # Encode the mermaid code
        encoded = encode_mermaid(mermaid_code)

        # Use mermaid.ink API to get the image
        url = f"https://mermaid.ink/img/{encoded}"

        # Download the image
        response = requests.get(url, timeout=30)
        response.raise_for_status()

        # Save to file
        with open(output_file, 'wb') as f:
            f.write(response.content)

        print(f"Generated flowchart image: {output_file}")
        return True
    except Exception as e:
        print(f"Error generating flowchart: {e}")
        return False

def parse_markdown_to_word(md_file, output_file):
    """Parse Markdown file and convert to Word document"""

    # Read markdown file
    print(f"Reading {md_file}...")
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create Word document
    doc = Document()

    # Set default font to support Chinese
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Split content into lines
    lines = content.split('\n')

    i = 0
    in_mermaid = False
    mermaid_code = []
    mermaid_index = 0

    while i < len(lines):
        line = lines[i]

        # Handle Mermaid code blocks
        if line.strip() == '```mermaid':
            in_mermaid = True
            mermaid_code = []
            i += 1
            continue

        if in_mermaid and line.strip() == '```':
            in_mermaid = False
            mermaid_text = '\n'.join(mermaid_code)

            # Generate image
            image_file = rf'd:\lasthm\gangzhu\temp_mermaid_{mermaid_index}.png'
            if mermaid_to_image(mermaid_text, image_file):
                try:
                    # Add image to document
                    doc.add_picture(image_file, width=Inches(6))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Error adding image to document: {e}")
                    # Fallback: add as code
                    p = doc.add_paragraph()
                    run = p.add_run(mermaid_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)

            mermaid_index += 1
            i += 1
            continue

        if in_mermaid:
            mermaid_code.append(line)
            i += 1
            continue

        # Handle headings
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=4)
            p.runs[0].font.color.rgb = RGBColor(149, 165, 166)
            i += 1
            continue

        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            p.runs[0].font.color.rgb = RGBColor(127, 140, 141)
            i += 1
            continue

        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            p.runs[0].font.color.rgb = RGBColor(52, 73, 94)
            i += 1
            continue

        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.runs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            i += 1
            continue

        # Handle horizontal rules
        if line.strip() == '---':
            doc.add_paragraph('_' * 80)
            i += 1
            continue

        # Handle bold text
        if '**' in line:
            line = line.replace('**', '')

        # Handle empty lines
        if not line.strip():
            i += 1
            continue

        # Handle list items
        if re.match(r'^\d+\.\s', line):
            # Numbered list
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        if line.strip().startswith('- '):
            # Bullet list
            text = line.strip()[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Handle code blocks (non-mermaid)
        if line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```

            # Add code as preformatted text
            p = doc.add_paragraph('\n'.join(code_lines))
            run = p.runs[0]
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(44, 62, 80)
            p.style = 'No Spacing'
            continue

        # Handle table rows (simplified)
        if '|' in line and not line.strip().startswith('|'):
            i += 1
            continue

        # Regular paragraph
        if line.strip():
            doc.add_paragraph(line.strip())

        i += 1

    # Save document
    print(f"Saving to {output_file}...")
    doc.save(output_file)
    print(f"Successfully created Word document: {output_file}")

if __name__ == '__main__':
    md_file = r'd:\lasthm\gangzhu\技术方案.md'
    output_file = r'd:\lasthm\gangzhu\技术方案.docx'

    try:
        parse_markdown_to_word(md_file, output_file)
        print("\nConversion completed!")
    except Exception as e:
        print(f"Error during conversion: {e}")
        import traceback
        traceback.print_exc()
