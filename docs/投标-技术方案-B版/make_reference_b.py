# -*- coding: utf-8 -*-
"""生成 C 版投标文件标准排版 reference.docx（pandoc --reference-doc 用）。
排版规范（投标行业标准）：
- 纸张 A4；页边距 上/下 2.54cm、左 3.0cm(装订)、右 2.5cm；页眉/页脚 1.75cm
- 正文：宋体小四(12pt)，西文 Times New Roman，黑色，固定行距 22 磅，
        两端对齐，首行缩进 2 字符，段前段后 0
- 一级标题(章)：黑体三号(16pt)加粗，居中
- 二级标题(节)：黑体四号(14pt)加粗，左对齐
- 三级标题：宋体小四(12pt)加粗，左对齐
- 四级标题：宋体小四(12pt)加粗
- 图/表题注：宋体五号(10.5pt)，居中
- 表格正文：仿宋五号(10.5pt)
- 页脚：页面底端居中阿拉伯数字页码
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH


def set_font(style, ascii_font, east_font, size_pt, bold=False, color=None):
    style.font.name = ascii_font
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), ascii_font)
    rfonts.set(qn('w:hAnsi'), ascii_font)
    rfonts.set(qn('w:eastAsia'), east_font)
    rfonts.set(qn('w:cs'), ascii_font)


def set_first_line_indent_chars(pf, chars=2):
    """首行缩进 N 字符（用 firstLineChars，随字号自适应）。"""
    ppr = pf.element.get_or_add_pPr()
    ind = ppr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        ppr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars * 100))
    ind.set(qn('w:firstLine'), str(int(chars * 240)))  # 兜底 twips


def exact_line(pf, pt=22):
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(pt)


doc = Document('reference-default.docx')
black = RGBColor(0, 0, 0)

# ---------- 正文：宋体小四，固定行距22磅，两端对齐，首行缩进2字符 ----------
for body_style in ['Normal', 'Body Text', 'First Paragraph']:
    try:
        st = doc.styles[body_style]
    except KeyError:
        continue
    set_font(st, 'Times New Roman', '宋体', 12, color=black)
    pf = st.paragraph_format
    exact_line(pf, 22)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_first_line_indent_chars(pf, 2)

# ---------- 标题分级 ----------
# Heading 1 章：黑体三号加粗居中
h1 = doc.styles['Heading 1']
set_font(h1, 'Times New Roman', '黑体', 16, bold=True, color=black)
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(12)
h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
# 一级标题清除首行缩进
_p = h1.paragraph_format.element.get_or_add_pPr()
_ind = _p.find(qn('w:ind'))
if _ind is not None:
    _p.remove(_ind)

# Heading 2 节：黑体四号加粗左对齐
h2 = doc.styles['Heading 2']
set_font(h2, 'Times New Roman', '黑体', 14, bold=True, color=black)
h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(6)
h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

# Heading 3：宋体小四加粗
h3 = doc.styles['Heading 3']
set_font(h3, 'Times New Roman', '宋体', 12, bold=True, color=black)
h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h3.paragraph_format.space_before = Pt(8)
h3.paragraph_format.space_after = Pt(4)
h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

# Heading 4：宋体小四加粗
try:
    h4 = doc.styles['Heading 4']
    set_font(h4, 'Times New Roman', '宋体', 12, bold=True, color=black)
    h4.paragraph_format.space_before = Pt(6)
    h4.paragraph_format.space_after = Pt(2)
    h4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
except KeyError:
    pass

# ---------- 图/表题注：宋体五号居中 ----------
for aux in ['Caption', 'Image Caption', 'Table Caption']:
    try:
        st = doc.styles[aux]
        set_font(st, 'Times New Roman', '宋体', 10.5, color=black)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    except KeyError:
        pass

# ---------- 表格正文：仿宋五号（pandoc 表格单元格用 Compact/Table 样式） ----------
for tbl_style in ['Compact', 'Table Paragraph']:
    try:
        st = doc.styles[tbl_style]
        set_font(st, 'Times New Roman', '仿宋', 10.5, color=black)
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        st.paragraph_format.space_before = Pt(1)
        st.paragraph_format.space_after = Pt(1)
        _pp = st.paragraph_format.element.get_or_add_pPr()
        _i = _pp.find(qn('w:ind'))
        if _i is not None:
            _pp.remove(_i)
    except KeyError:
        pass

# ---------- 页面：A4 + 页边距 + 页眉页脚距离 ----------
for sec in doc.sections:
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    sec.header_distance = Cm(1.75)
    sec.footer_distance = Cm(1.75)

# ---------- 页脚：居中阿拉伯数字页码 ----------
for sec in doc.sections:
    footer = sec.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # 宋体五号页码
    rpr = run._element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Times New Roman')
    rf.set(qn('w:hAnsi'), 'Times New Roman')
    rf.set(qn('w:eastAsia'), '宋体')
    rpr.append(rf)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')  # 10.5pt*2
    rpr.append(sz)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._element.append(fld)

doc.save('reference.docx')
print('reference.docx OK')
